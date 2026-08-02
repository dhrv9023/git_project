from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    doc.add_paragraph(text)

def bullet(text, level=0):
    doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')

def warning(text):
    p = doc.add_paragraph()
    run = p.add_run("⚠ " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light List Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    doc.add_paragraph()

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading('StockBuddy — Senior Quant/ML Engineer Portfolio Review', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Reviewer posture: Hiring manager at a quant fund / ML-heavy fintech. Optimism is not the goal. Accuracy is.')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
doc.add_paragraph()

# ── Section 1 ────────────────────────────────────────────────────────────────
h1('1. Overall Ratings')
add_table(
    ['Axis', 'Score', 'Justification'],
    [
        ['Technical Depth / ML-Quant Rigor', '4 / 10',
         'Correct vocabulary, but no walk-forward validation, no IC, no statistical test for regime predictiveness.'],
        ['Code Architecture & Engineering', '4 / 10',
         '1100-line God-function file, global mutable CONFIG, no tests, no logging framework, retraining per-request is an anti-pattern.'],
        ['Real-World Usability', '2 / 10',
         'Minutes-long training per request, no auth, no rate-limiting, no mobile layout, zero compliance disclaimers.'],
        ['Resume / Portfolio Impact', '5 / 10',
         'Impressive breadth for a solo project; a technical quant interviewer finds the cracks in under 5 minutes.'],
        ['Originality', '5 / 10',
         'Regime clustering + cosine scenario matching is a non-trivial concept; the unvalidated execution limits the score.'],
    ]
)

# ── Section 2 ────────────────────────────────────────────────────────────────
h1('2. Critical Flaws')

h2('2.1  On-Demand Training With No Persistence — Red Flag')
body('Every call to /api/predict trains LSTM, GRU, and Transformer from scratch. On CPU this takes 5–15 minutes. '
     'There is no model.save(), no versioning, no caching. A hiring manager who asks "how do you ensure the model you '
     'show today isn\'t worse than yesterday\'s?" has no good answer to receive from this codebase.')
warning('Absence of model persistence signals "made it work once" thinking, not production engineering.')

h2('2.2  K-Means Regime Labeling — Statistically Cosmetic')
body('Clusters are reordered post-hoc using a heuristic score: 2*Ret10 + 1.5*RSI_dev - 3*Vol. '
     'This gives human-readable names ("Trending Bull", "Bear Stress") but the labels are never validated against forward returns.')
bullet('No p-value or significance test on whether Regime 0 predicts higher forward returns than Regime 4.')
bullet('No elbow curve, silhouette score, or stability analysis justifying k=6.')
bullet('compute_regime_stats() computes forward return tables but never tests for statistical significance.')
warning('Asserting regime labels are predictive without proof is data-mining — a quant interviewer will call this out immediately.')

h2('2.3  Backtest — Lookahead Bias Present')
body('The K-Means model is fit on the ENTIRE historical dataset — including the "test" period — before the backtest runs. '
     'Regime labels assigned to day T use cluster centers computed from data up to T+N. This is a classic, disqualifying backtest error.')
bullet('Fix: Fit K-Means only on training window; assign regimes on held-out data by projecting onto FROZEN cluster centers.')
bullet('No transaction costs or slippage modeled — daily regime rebalancing with zero friction is unrealistic.')
bullet('Test set ≈ 112 samples for models with thousands of parameters — reported metrics are optimistically biased.')
warning('Lookahead bias in a backtest is the single most disqualifying error for a quant audience. Fix this before showing the project to any quant interviewer.')

h2('2.4  Other Assumptions That Fall Apart Under Scrutiny')
add_table(
    ['Assumption', 'Problem'],
    [
        ['yfinance as data source', 'Unofficial scraper; known data gaps and split-adjustment errors. Production tools use Polygon.io, Refinitiv, etc.'],
        ['MinMaxScaler fit on full dataset', 'Scaler sees future test data during fit — leaks future scale info into training. Scalers must be fit on train set only.'],
        ['Confidence intervals via model spread', '±1.96 × std of ensemble disagreement is not a calibrated confidence interval. Calling it one misleads users.'],
        ['Single-ticker analysis', 'Regime analysis is most powerful in a portfolio context. Single-ticker regime does not generalize.'],
        ['Port inconsistency', 'Code runs on port 5000; README and summary say 8080. Signals docs are not kept in sync with code.'],
        ['No requirements.txt', 'A project with no dependency file is not reproducible. Most basic signal of engineering maturity.'],
    ]
)

# ── Section 3 ────────────────────────────────────────────────────────────────
h1('3. Enhancement Roadmap')

h2('🔴 Must-Fix Before Showing to Employers')
for item in [
    'Fix lookahead bias: fit K-Means and scalers only on training data; project onto held-out splits using frozen parameters.',
    'Fix MinMaxScaler leakage: fit on train, transform val/test with the train-fitted scaler.',
    'Add model persistence: model.save() keyed by (ticker, date_range, config_hash). Load if exists, train if not.',
    'Add statistical validation of regime labels: Kruskal-Wallis or ANOVA on forward 10-day returns across regime groups.',
    'Add requirements.txt or pyproject.toml.',
    'Fix port inconsistency between code (5000) and documentation (8080).',
]:
    bullet(item)

h2('🟡 High-Impact Additions')
for item in [
    'Walk-Forward Cross-Validation: replace single 70/15/15 split with expanding/rolling window — industry standard for time-series ML.',
    'Information Coefficient (IC): Spearman correlation between predicted and realized log-returns. IC > 0.05 is meaningful in quant finance.',
    'Transaction-cost-aware backtest: add commission_bps and slippage_bps parameters.',
    'Async training: return a job_id immediately, poll for completion — makes UX actually usable.',
    'Multi-ticker comparison: watchlist of 5–10 tickers with a regime heatmap grid.',
    'Proper uncertainty quantification: conformal prediction or Monte Carlo Dropout instead of model-spread proxy.',
]:
    bullet(item)

h2('🟢 Nice-to-Have Polish')
for item in [
    'Unit tests for compute_rsi, compute_macd, engineer_features, run_backtest.',
    'Structured logging (Python logging module) instead of bare print() statements.',
    'Complete type annotations throughout app.py.',
    'Docker + docker-compose for one-command deployment.',
    'Mobile-responsive CSS for the frontend.',
    'Hosted live demo (Railway / Render / HuggingFace Spaces) — portfolio projects need a live URL.',
]:
    bullet(item)

# ── Section 4 ────────────────────────────────────────────────────────────────
h1('4. Real-World Usability Gaps')
add_table(
    ['Gap', 'Severity', 'Notes'],
    [
        ['Training latency (5–15 min)', '🔴 Critical', 'Renders AI features unusable for real users.'],
        ['No authentication', '🔴 Critical', 'Anyone can hammer the training endpoint and exhaust compute.'],
        ['No financial disclaimer', '🔴 Legal', 'Labels like "Favorable Conditions" create financial advice liability.'],
        ['No rate limiting', '🔴 Critical', 'Single malicious request can DoS the server.'],
        ['No mobile layout', '🟡 High', 'Fixed-width Chart.js containers are unusable on phones.'],
        ['No deployment story', '🟡 High', 'No Dockerfile, no WSGI wrapper — flask dev server is not production.'],
        ['No watchlist/persistence', '🟡 Medium', 'Users cannot save analyses or compare across sessions.'],
        ['No data freshness indicator', '🟡 Medium', 'Users cannot tell if data is current or stale (weekend, holiday).'],
        ['No accessibility (a11y)', '🟢 Low', 'Color-coded charts with no alt-text or accessible labels.'],
    ]
)

# ── Section 5 ────────────────────────────────────────────────────────────────
h1('5. Differentiation Analysis')

h2('vs. TradingView / Yahoo Finance')
body('TradingView has RSI, MACD, EMA, and a library of community indicators natively. StockBuddy does not compete on indicator breadth. '
     'The regime clustering + cosine scenario matching is genuinely not exposed in either tool as an interactive visual. '
     'That angle is novel as a UX concept — but only if the statistics are defensible (see §2.2).')

h2('vs. Typical Bootcamp Stock Predictors')
body('The typical bootcamp project: download AAPL, train LSTM on Close prices, show it "predicts" the price by lagging yesterday\'s value. '
     'StockBuddy is meaningfully more sophisticated: targets log-returns not prices, engineers proper technical features, '
     'ensembles three architectures, and adds regime analysis. A genuine step up — but the backtest flaws pull it back toward tutorial territory for a quant audience.')

h2('vs. QuantConnect / Alpaca')
body('These are professional platforms with event-driven simulation, survivorship-bias-corrected data, slippage models, and execution integrations. '
     'StockBuddy is NOT in this category. Correct framing: StockBuddy is a research visualization tool, not a trading system.')

h2('The Regime + Cosine Matching Angle — Genuinely Differentiated?')
body('Partially yes. K-Means market regimes are used in real quant research (GS, JPMorgan publish regime-overlay papers). '
     'Cosine-similarity scenario matching is a legitimate technique. The combination in an interactive visual is a non-trivial product concept. '
     'Score is limited because the implementation is not statistically validated — the tool asserts its regimes are meaningful without proving it.')

# ── Section 6 ────────────────────────────────────────────────────────────────
h1('6. Resume Framing')
body('Calibrated to the ACTUAL sophistication level — not inflated.')

h2('Option A — ML / Data Science Roles (recommended)')
p = doc.add_paragraph()
p.add_run('Built a quantitative market intelligence dashboard using Python (Flask), TensorFlow, and scikit-learn; engineered an unsupervised '
          'K-Means market regime classifier (k=6) over RSI/MACD/volatility feature vectors, implemented a cosine-similarity historical scenario '
          'matching engine, and ensembled LSTM/GRU/Transformer models to generate directional return signals with regime-based backtesting.').italic = True

h2('Option B — Software Engineering Roles')
p = doc.add_paragraph()
p.add_run('Developed a full-stack financial analytics platform (Flask REST API + vanilla JS SPA) featuring live yfinance data ingestion, '
          'interactive Chart.js visualizations, and on-demand TensorFlow model training; designed a 3-endpoint RESTful API with JSON-safe '
          'NaN/Inf handling and MinMaxScaler feature pipelines.').italic = True

h2('Option C — Quant Analyst / Researcher (only after fixing §2.1–2.3)')
p = doc.add_paragraph()
p.add_run('Designed and backtested a regime-overlay tactical allocation strategy using unsupervised market state classification and '
          'cosine-similarity scenario matching on 15+ years of daily technical indicators; implemented an LSTM/GRU/Transformer ensemble '
          'with walk-forward validation and transaction-cost-aware backtest reporting.').italic = True
warning('Do NOT use Option C until lookahead bias and data leakage issues are fixed. A quant interviewer WILL ask you to walk through the methodology.')

# ── Section 7 ────────────────────────────────────────────────────────────────
h1('7. What to Add to Make StockBuddy Usable for Everyone')

h2('For the General Retail / Non-Technical User')
add_table(
    ['Feature', 'Why It Matters'],
    [
        ['Plain-English Regime Summary', 'Replace "Regime 0: Trending Bull" with "📈 Stocks in this condition gained +3.2% on average over the next 10 days." Most users don\'t know what RSI means.'],
        ['"What Should I Do?" Panel', 'Users want a simple takeaway. If GREEN + low risk → "Historically favorable to hold. NOT financial advice."'],
        ['Watchlist / Portfolio Mode', 'Let users track 5–10 tickers simultaneously with a regime heatmap grid. The #1 missing usability feature.'],
        ['Email/Telegram Regime Alerts', 'Users want notifications when a regime changes, not to manually re-check a dashboard every day.'],
        ['Mobile-First Responsive Design', 'Most retail users access tools on phones. Current dashboard is desktop-only.'],
        ['Onboarding Tooltips / Guided Tour', 'First-time users need a 30-second explanation of RSI, regime, and risk score.'],
        ['Historical Accuracy Scorecard', '"When the dashboard said GREEN in the last 12 months, the stock went up X% of the time." Builds trust through transparency.'],
    ]
)

h2('For the Technical / Quant User')
add_table(
    ['Feature', 'Why It Matters'],
    [
        ['Ticker Comparison Mode', 'Show two tickers\' regimes side-by-side. Useful for pair-trading research.'],
        ['Custom Regime Training Window', 'A regime model trained on 2010–2018 may not apply to post-COVID market structure.'],
        ['CSV/JSON Data Export', 'Let users take regime timeline and signal data into their own tools (Excel, R, Python).'],
        ['Walk-Forward Backtest Chart', 'Show how the strategy\'s performance evolved over time, revealing if the edge is stable or decaying.'],
        ['Strategy Parameter Sweep', 'Let users adjust allocation weights (100/50/0%) and replay the backtest — builds intuition.'],
        ['Factor Exposure Overlay', 'Show which factor exposures (momentum, vol, value) characterize each regime.'],
    ]
)

h2('For Deployment / Productionization')
add_table(
    ['Feature', 'Why It Matters'],
    [
        ['Docker + docker-compose', 'One-command deployable. Necessary for anyone to actually run this.'],
        ['Hosted Live Demo', 'Portfolio projects need a live URL. "Clone and run locally" is not a demo.'],
        ['User Accounts + Watchlist Persistence', 'Users won\'t return to a tool that forgets their settings (SQLite or Supabase).'],
        ['Pre-trained Model Cache', 'Ship pre-trained models for top 20 tickers (AAPL, TSLA, MSFT, etc.) so AI features work instantly.'],
        ['Financial Disclaimer Footer', 'Legally required. One line: "For educational purposes only. Not financial advice."'],
        ['API Rate Limiting (Flask-Limiter)', 'Prevents abuse of the training endpoint.'],
    ]
)

# ── Verdict ───────────────────────────────────────────────────────────────────
h1('Final Verdict')
doc.add_paragraph(
    'StockBuddy is not resume-ready for quant or ML engineering roles as-is, but it is resume-ready for software engineering '
    'and data science roles at non-quant companies. The project demonstrates genuine breadth — a full-stack implementation with '
    'real ML techniques, thoughtful UX, and a non-trivial quant concept — which puts it well above the average portfolio project. '
    'However, it contains three disqualifying errors for any rigorous technical audience: (1) lookahead bias in the backtest, '
    '(2) MinMaxScaler fit on full data including the test set, and (3) regime labels that are asserted to be predictive but never '
    'statistically validated. Any quant interviewer who asks "walk me through your backtest methodology" will find issue (1) within '
    'two minutes.'
)
warning(
    'The one thing that would change the answer: Fix the lookahead bias, add an IC or regime-conditional return significance test, '
    'and deploy a live hosted demo. With those three changes, this becomes a portfolio piece that can credibly claim quant-adjacent '
    'ML research — a rare and genuinely impressive category for a solo project.'
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/home/dhruv/Desktop/stockbuddy/StockBuddy/StockBuddy_Portfolio_Review.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
